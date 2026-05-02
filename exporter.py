import re
import tempfile
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF



def _clean(text):
    return re.sub(r'[*_`#]', '', text).strip()


def _extract_abstract(content):
    lines, in_abs, buf = content.split('\n'), False, []
    for line in lines:
        s = line.strip()
        if re.search(r'abstract', s, re.IGNORECASE) and s.startswith('#'):
            in_abs = True
            continue
        if in_abs:
            if s.startswith('#') or re.match(r'^[IVX]+\.', s):
                break
            if s:
                buf.append(_clean(s))
    return ' '.join(buf)


def _add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _set_two_columns(section):
    sectPr = section._sectPr
    existing = sectPr.find(qn('w:cols'))
    if existing is not None:
        sectPr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)


def _add_page_number(footer_paragraph):
    run = footer_paragraph.add_run()
    run.font.size = Pt(8)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _buf_to_tempfile(buf):
    """Save a BytesIO PNG buffer to a temp file, return path."""
    buf.seek(0)
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        tmp.write(buf.read())
        return tmp.name


def _insert_docx_image(doc, buf, caption, width=Inches(5.5)):
    """Insert a BytesIO image into a docx document with a caption."""
    try:
        tmp_path = _buf_to_tempfile(buf)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(tmp_path, width=width)
        os.unlink(tmp_path)
    
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        r = cap.add_run(caption)
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(9)
        r.font.italic = True
    except Exception as e:
        print(f"DOCX image insert skipped ({caption}): {e}")



ROMAN = ['I','II','III','IV','V','VI','VII','VIII','IX','X']


def save_ieee_docx(topic, content, authors="Research Agent", images=None):
    """
    Generate IEEE-formatted .docx with:
    - Single-column title + abstract
    - Two-column body
    - Figures embedded after relevant sections
    - Correct Roman numeral headings (A./B./C. subsections, not Roman)
    - No abstract duplication
    """
    filename = topic.replace(" ", "_") + "_IEEE_paper.docx"
    doc      = Document()
    images   = images or {}

    sec0 = doc.sections[0]
    sec0.page_width    = Inches(8.5)
    sec0.page_height   = Inches(11)
    sec0.left_margin   = Inches(0.65)
    sec0.right_margin  = Inches(0.65)
    sec0.top_margin    = Inches(0.75)
    sec0.bottom_margin = Inches(0.75)

    hp = sec0.header.paragraphs[0]
    hp.text      = "IEEE Open Journal — Auto-Generated Research Paper"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.runs[0].font.name      = 'Times New Roman'
    hp.runs[0].font.size      = Pt(8)
    hp.runs[0].font.italic    = True
    hp.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    fp = sec0.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(fp)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(topic)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(authors)
    r.font.name   = 'Times New Roman'
    r.font.size   = Pt(10)
    r.font.italic = True

    _add_rule(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("Abstract")
    r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    r.font.bold = True; r.font.italic = True

    abstract = _extract_abstract(content) or f"This paper presents a study on {topic}."
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(abstract)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)

    _add_rule(doc)

    body_sec = doc.add_section(WD_SECTION.CONTINUOUS)
    body_sec.page_width    = Inches(8.5)
    body_sec.page_height   = Inches(11)
    body_sec.left_margin   = Inches(0.65)
    body_sec.right_margin  = Inches(0.65)
    body_sec.top_margin    = Inches(0.75)
    body_sec.bottom_margin = Inches(0.75)
    _set_two_columns(body_sec)

    sec_idx       = 0
    in_code       = False
    code_buf      = []
    skip_abstract = False

    def _body_text(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after       = Pt(3)
        p.paragraph_format.first_line_indent = Inches(0.15)
        r = p.add_run(_clean(text))
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)

    def _section_head(text):
        nonlocal sec_idx
        label = ROMAN[sec_idx] if sec_idx < len(ROMAN) else str(sec_idx + 1)
        sec_idx += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(f"{label}. {_clean(text).upper()}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.font.bold = True

        
        low = text.lower()
        if "introduction" in low and images.get("architecture"):
            _insert_docx_image(doc, images["architecture"], "Fig. 1: System Architecture Diagram")
        elif ("method" in low or "proposed" in low) and images.get("flowchart"):
            _insert_docx_image(doc, images["flowchart"], "Fig. 2: Methodology Flowchart")
        elif "result" in low or "experiment" in low:
            if images.get("comparison"):
                _insert_docx_image(doc, images["comparison"], "Fig. 3: Performance Comparison")
            if images.get("training"):
                _insert_docx_image(doc, images["training"], "Fig. 4: Training Curves")
        elif "related" in low or "literature" in low:
            if images.get("graph"):
                _insert_docx_image(doc, images["graph"], "Fig. 5: Citation Network")
            if images.get("year_chart"):
                _insert_docx_image(doc, images["year_chart"], "Fig. 6: Papers by Year")

    def _subsection_head(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(_clean(text))
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(10)
        r.font.bold   = True
        r.font.italic = True

    def _code_block(lines):
        for cl in lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(0)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),  'clear')
            shd.set(qn('w:fill'), 'F0F0F0')
            pPr.append(shd)
            r = p.add_run(cl.replace('\t', '    '))
            r.font.name = 'Courier New'
            r.font.size = Pt(8)
        doc.add_paragraph()

    for line in content.split('\n'):
        stripped = line.strip()

        if stripped.startswith('```'):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False; _code_block(code_buf)
            continue

        if in_code:
            code_buf.append(line)
            continue

        if not stripped:
            continue

        if re.search(r'^#{1,3}\s*abstract', stripped, re.IGNORECASE):
            skip_abstract = True
            continue

        if skip_abstract:
            if re.match(r'^#{1,3}\s', stripped) or re.match(r'^[IVX]+\.\s', stripped):
                skip_abstract = False 
            else:
                continue

        if stripped.lower() in ('see above.', 'see above'):
            continue

        if re.match(r'^[A-Z]\.\s', stripped):
            _subsection_head(stripped)

        elif re.match(r'^#{1,3}\s', stripped) or re.match(r'^[IVX]+\.\s', stripped):
            title = re.sub(r'^#{1,3}\s', '', stripped)
            title = re.sub(r'^[IVX]+\.\s', '', title)
            _section_head(title)

        else:
            _body_text(stripped)

    doc.save(filename)
    print(f"IEEE DOCX saved: {filename}")
    return filename


class UnicodePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font("DejaVu",  "",  "DejaVuSans.ttf",         uni=True)
        self.add_font("DejaVu",  "B", "DejaVuSans-Bold.ttf",    uni=True)
        self.add_font("DejaVu",  "I", "DejaVuSans-Oblique.ttf", uni=True)

    def header(self):
        self.set_font("DejaVu", "I", 8)
        self.set_text_color(130, 130, 130)
        self.cell(0, 8, "IEEE Style Research Paper — Auto Generated", align="C", ln=True)
        self.set_draw_color(180, 180, 180)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(2)

    def footer(self):
        self.set_y(-13)
        self.set_font("DejaVu", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def save_pdf(topic, content, images=None):
    filename = topic.replace(" ", "_") + "_paper.pdf"
    pdf = UnicodePDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    images = images or {}

    pdf.set_font("DejaVu", "B", 16)
    pdf.set_fill_color(10, 10, 60)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 14, topic, fill=True, ln=True, align="C")
    pdf.ln(4)
    pdf.set_text_color(0, 0, 0)

    in_code = False; code_buf = []; skip_abstract = False

    def flush_code(lines):
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font("Courier", "", 8)
        pdf.set_text_color(30, 30, 30)
        for cl in lines:
            pdf.multi_cell(0, 5, cl.replace('\t', '    '), fill=True)
        pdf.ln(2)

    for line in content.split('\n'):
        stripped = line.strip()

        if stripped.startswith('```'):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False; flush_code(code_buf)
            continue

        if in_code:
            code_buf.append(line)
            continue

        if not stripped:
            pdf.ln(2)
            continue

        if re.search(r'^#{1,3}\s*abstract', stripped, re.IGNORECASE):
            skip_abstract = True
            continue

        if skip_abstract:
            if re.match(r'^#{1,3}\s', stripped) or re.match(r'^[IVX]+\.\s', stripped):
                skip_abstract = False
            else:
                continue

        if re.match(r'^[A-Z]\.\s', stripped):
            pdf.set_font("DejaVu", "B", 10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, _clean(stripped))
            pdf.ln(1)

        elif re.match(r'^#{1,3} ', stripped) or re.match(r'^[IVX]+\.', stripped):
            heading = re.sub(r'^#{1,3} ', '', stripped).replace('**', '').strip()
            pdf.set_font("DejaVu", "B", 12)
            pdf.set_text_color(10, 10, 60)
            pdf.set_fill_color(220, 225, 255)
            pdf.cell(0, 8, _clean(heading), fill=True, ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)

            low = heading.lower()
            if "introduction" in low and images.get("architecture"):
                _insert_pdf_image(pdf, images["architecture"], "Fig. 1: System Architecture")
            elif ("method" in low or "proposed" in low) and images.get("flowchart"):
                _insert_pdf_image(pdf, images["flowchart"], "Fig. 2: Methodology Flowchart")
            elif "result" in low or "experiment" in low:
                if images.get("comparison"):
                    _insert_pdf_image(pdf, images["comparison"], "Fig. 3: Performance Comparison")
                if images.get("training"):
                    _insert_pdf_image(pdf, images["training"], "Fig. 4: Training Curves")
            elif "related" in low or "literature" in low:
                if images.get("graph"):
                    _insert_pdf_image(pdf, images["graph"], "Fig. 5: Citation Network")
                if images.get("year_chart"):
                    _insert_pdf_image(pdf, images["year_chart"], "Fig. 6: Papers by Year")

        elif "**" in stripped:
            clean2 = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            pdf.set_font("DejaVu", "B", 10)
            pdf.multi_cell(0, 6, re.sub(r'[*_`#]', '', clean2))

        else:
            pdf.set_font("DejaVu", "", 10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, re.sub(r'[*_`#]', '', stripped))

    pdf.output(filename)
    print(f"PDF saved: {filename}")
    return filename


def _insert_pdf_image(pdf, buf, caption):
    try:
        tmp_path = _buf_to_tempfile(buf)
        pdf.ln(2)
        pdf.image(tmp_path, x=20, w=170)
        pdf.set_font("DejaVu", "I", 8)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 6, caption, align="C", ln=True)
        pdf.ln(3)
        os.unlink(tmp_path)
    except Exception as e:
        print(f"PDF image insert skipped ({caption}): {e}")


def save_markdown(topic, content):
    filename = topic.replace(" ", "_") + "_paper.md"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"💾 Markdown saved: {filename}")
    return filename
